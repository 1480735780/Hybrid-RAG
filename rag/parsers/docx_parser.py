"""
DOCX Parser - Extract text from Word documents.
"""

from typing import Dict

from loguru import logger

from rag.parsers.document_parser import BaseParser


class DocxParser(BaseParser):
    """
    DOCX Parser for extracting text from Word documents.

    Features:
    - Paragraph extraction
    - Table extraction
    - Metadata extraction
    """

    def parse(self, file_path: str) -> str:
        """
        Parse DOCX and return text content.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text content
        """
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        text_parts.append(" | ".join(row_text))

            return "\n\n".join(text_parts)

        except ImportError:
            logger.error("python-docx not installed")
            raise
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise

    def parse_with_metadata(self, file_path: str) -> Dict:
        """
        Parse DOCX and return text with metadata.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Dict with text and metadata
        """
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []
            paragraphs = []
            tables = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name if para.style else None,
                    })

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        table_data.append(row_data)
                        text_parts.append(" | ".join(row_data))

                if table_data:
                    tables.append({
                        "table_index": table_idx,
                        "rows": len(table_data),
                        "data": table_data,
                    })

            # Extract core properties
            metadata = {
                "file_path": file_path,
                "file_type": "docx",
                "total_paragraphs": len(paragraphs),
                "total_tables": len(tables),
                "paragraphs": paragraphs,
                "tables": tables,
            }

            # Add document properties if available
            if doc.core_properties:
                props = doc.core_properties
                metadata.update({
                    "title": props.title or "",
                    "author": props.author or "",
                    "subject": props.subject or "",
                    "created": str(props.created) if props.created else None,
                    "modified": str(props.modified) if props.modified else None,
                })

            return {
                "text": "\n\n".join(text_parts),
                "metadata": metadata,
            }

        except Exception as e:
            logger.error(f"Error parsing DOCX with metadata: {e}")
            raise
